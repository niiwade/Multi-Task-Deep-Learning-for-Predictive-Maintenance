#!/usr/bin/env python3
"""
Simple markdown to DOCX converter - creates professional formatted DOCX
"""

import sys

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

def simple_md_to_docx(md_file, docx_file):
    """Convert markdown to DOCX with basic formatting"""
    
    doc = Document()
    
    # Add title page
    title = doc.add_heading('Multi-Task Deep Learning for Predictive Maintenance', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Final Project Report')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(14)
        run.font.bold = True
    
    date_para = doc.add_paragraph('January 30, 2026')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    skip_next = False
    
    for i, line in enumerate(lines):
        if skip_next:
            skip_next = False
            continue
            
        stripped = line.strip()
        
        # Skip empty lines and title
        if not stripped or stripped.startswith('# Multi-Task'):
            continue
        
        # Heading 2
        if stripped.startswith('## '):
            text = stripped[3:].strip()
            if text and text != '---':
                heading = doc.add_heading(text, level=1)
            continue
        
        # Heading 3
        if stripped.startswith('### '):
            text = stripped[4:].strip()
            heading = doc.add_heading(text, level=2)
            continue
        
        # Heading 4
        if stripped.startswith('#### '):
            text = stripped[5:].strip()
            heading = doc.add_heading(text, level=3)
            continue
        
        # Horizontal rule
        if stripped == '---':
            doc.add_paragraph()
            continue
        
        # Bullet points
        if stripped.startswith('- '):
            text = stripped[2:].strip()
            doc.add_paragraph(text, style='List Bullet')
            continue
        
        # Code block
        if stripped.startswith('```'):
            code_lines = []
            j = i + 1
            while j < len(lines) and not lines[j].strip().startswith('```'):
                code_lines.append(lines[j])
                j += 1
            
            code_text = '\n'.join(code_lines).strip()
            if code_text:
                p = doc.add_paragraph(code_text)
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                p.paragraph_format.left_indent = Inches(0.3)
            continue
        
        # Tables (simple pipe-separated)
        if '|' in stripped and not stripped.startswith('###'):
            # Check if this is table header
            table_data = []
            j = i
            
            # Collect table rows
            while j < len(lines):
                row = lines[j].strip()
                if not row or not '|' in row:
                    break
                if '---' not in row:  # Skip separator rows
                    cells = [c.strip() for c in row.split('|')[1:-1]]
                    if cells:
                        table_data.append(cells)
                j += 1
            
            if table_data and len(table_data) > 0:
                # Create table
                cols = len(table_data[0])
                rows = len(table_data)
                
                table = doc.add_table(rows=rows, cols=cols)
                table.style = 'Light Grid Accent 1'
                
                for row_idx, row_data in enumerate(table_data):
                    for col_idx, cell_text in enumerate(row_data):
                        try:
                            cell = table.rows[row_idx].cells[col_idx]
                            cell.text = cell_text
                            
                            # Style header row
                            if row_idx == 0 and '---' in (lines[i+1] if i+1 < len(lines) else ''):
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.bold = True
                        except:
                            pass
            
            skip_next = False
            continue
        
        # Regular paragraph
        if stripped:
            # Process basic markdown formatting
            text = stripped
            p = doc.add_paragraph()
            
            # Simple bold/italic replacement
            import re
            
            # Find **bold** and *italic*
            last_end = 0
            for match in re.finditer(r'\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`', text):
                # Add text before match
                if match.start() > last_end:
                    p.add_run(text[last_end:match.start()])
                
                # Add formatted text
                if match.group(1):  # Bold
                    run = p.add_run(match.group(1))
                    run.bold = True
                elif match.group(2):  # Italic
                    run = p.add_run(match.group(2))
                    run.italic = True
                elif match.group(3):  # Code
                    run = p.add_run(match.group(3))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                
                last_end = match.end()
            
            # Add remaining text
            if last_end < len(text):
                p.add_run(text[last_end:])
            
            # If nothing was added (no formatting), add as plain
            if not p.runs:
                p.text = stripped
    
    # Save document
    doc.save(docx_file)
    print(f"✅ Successfully created: {docx_file}")

if __name__ == '__main__':
    simple_md_to_docx('finalreport.md', 'finalreport.docx')
