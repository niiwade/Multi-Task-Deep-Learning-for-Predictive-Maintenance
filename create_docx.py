#!/usr/bin/env python3
"""
Convert finalreport.md to professional DOCX format
Uses pandoc or fallback to python-docx
"""

import os
import subprocess
import sys

def try_pandoc():
    """Try using pandoc to convert (best quality)"""
    try:
        result = subprocess.run(
            ['pandoc', 'finalreport.md', '-o', 'finalreport.docx'],
            cwd=os.getcwd(),
            capture_output=True,
            timeout=30
        )
        if result.returncode == 0:
            return True
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    return False

def create_with_docx():
    """Fallback: Create DOCX using python-docx"""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("Installing python-docx...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Read markdown
    with open('finalreport.md', 'r', encoding='utf-8', errors='ignore') as f:
        content = f.read()
    
    lines = content.split('\n')
    
    for line in lines:
        stripped = line.rstrip()
        
        # Title
        if stripped.startswith('# ') and not stripped.startswith('## '):
            title = stripped[2:].strip()
            if title and 'Multi-Task' in title:
                h = doc.add_heading(title, 0)
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        
        # Heading 2
        if stripped.startswith('## '):
            doc.add_heading(stripped[3:].strip(), level=1)
            continue
        
        # Heading 3
        if stripped.startswith('### '):
            doc.add_heading(stripped[4:].strip(), level=2)
            continue
        
        # Heading 4
        if stripped.startswith('#### '):
            doc.add_heading(stripped[5:].strip(), level=3)
            continue
        
        # Horizontal line
        if stripped == '---':
            doc.add_paragraph()
            continue
        
        # Bullet
        if stripped.startswith('- '):
            doc.add_paragraph(stripped[2:].strip(), style='List Bullet')
            continue
        
        # Code block
        if stripped.startswith('```'):
            continue
        
        # Table
        if '|' in stripped:
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            if cells and len(cells) > 0:
                doc.add_paragraph(' | '.join(cells))
            continue
        
        # Regular text
        if stripped and not stripped.startswith('**'):
            doc.add_paragraph(stripped)
    
    doc.save('finalreport.docx')
    return True

def main():
    print("Converting finalreport.md to DOCX...")
    print()
    
    # Try pandoc first (best output)
    print("[1/2] Attempting conversion with pandoc...")
    if try_pandoc():
        print("✅ Success with pandoc!")
        print("Created: finalreport.docx")
        return True
    
    # Fallback to python-docx
    print("[2/2] Pandoc not found, using python-docx...")
    if create_with_docx():
        print("✅ Success with python-docx!")
        print("Created: finalreport.docx")
        return True
    
    print("❌ Conversion failed")
    return False

if __name__ == '__main__':
    success = main()
    sys.exit(0 if success else 1)
