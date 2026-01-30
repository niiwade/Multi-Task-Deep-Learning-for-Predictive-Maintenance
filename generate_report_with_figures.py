"""
Convert finalreport.md to DOCX with embedded visualizations
Requires: python-docx library
Installation: pip install python-docx

Usage:
    python generate_report_with_figures.py
    
Output:
    finalreport_with_figures.docx (with embedded PNG images)
"""

from pathlib import Path
import os

def install_dependencies():
    """Install required packages"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        print("✓ python-docx already installed")
    except ImportError:
        print("📦 Installing python-docx...")
        os.system('pip install python-docx')
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    return Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH


def read_markdown_report(filepath):
    """Read the markdown report"""
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()


def create_docx_with_figures():
    """Create DOCX document with embedded figures"""
    
    print("\n" + "=" * 70)
    print("CREATING REPORT WITH EMBEDDED FIGURES")
    print("=" * 70)
    
    # Import dependencies
    Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH = install_dependencies()
    
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # ========================================================================
    # TITLE PAGE
    # ========================================================================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('Multi-Task Deep Learning for Predictive Maintenance')
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 102, 204)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run('Final Project Report with Visualizations')
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.italic = True
    
    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date.add_run('January 2026')
    date_run.font.size = Pt(12)
    
    doc.add_paragraph()  # Spacing
    
    # ========================================================================
    # TABLE OF CONTENTS / OVERVIEW
    # ========================================================================
    toc_heading = doc.add_heading('Report Overview', level=1)
    toc_heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)
    
    toc_items = [
        "Executive Q&A: All 6 project questions answered",
        "Topic Overview: Multi-task learning for industrial maintenance",
        "Current Knowledge: ML and TCN fundamentals",
        "Relevance: Business impact and safety implications",
        "Data Description: AI4I 2020 dataset specifications",
        "Methods & Tools: Architecture and implementation details",
        "Results (REAL): Phase 1-4 actual performance metrics",
        "Limitations: 8 critical root causes identified",
        "Recommendations: Actionable next steps",
        "Reflections & Lessons: Project insights"
    ]
    
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # ========================================================================
    # EXECUTIVE SUMMARY
    # ========================================================================
    exec_heading = doc.add_heading('Executive Summary', level=1)
    exec_heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)
    
    exec_summary = """
This project applies multi-task deep learning (Temporal Convolutional Networks with attention) 
to industrial predictive maintenance on the AI4I 2020 dataset. The model simultaneously performs:
    • Binary failure prediction (will the machine fail?)
    • Failure type classification (which failure mode?)
    • Time-to-failure regression (when will it fail?)

ACTUAL RESULTS vs TARGETS:
    • Binary Failure F1: 0.5545 (Target >0.95) ❌ -37% gap
    • Failure Type Accuracy: 97.25% (Target >95%) ✅ +2.3%
    • Time-to-Failure MAE: 3.50 hours (Target <2.0) ❌ +75% gap
    
KEY INSIGHT: Data quality, not model design, limited performance.
The 97:3 class imbalance with only 43 test failures makes the statistical 
confidence interval ±0.15 – targets were unrealistic from the start.
    """
    
    doc.add_paragraph(exec_summary)
    
    # ========================================================================
    # VISUALIZATION SECTION
    # ========================================================================
    viz_heading = doc.add_heading('Critical Visualizations', level=1)
    viz_heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)
    
    visualizations = [
        {
            'name': 'confusion_matrix.png',
            'title': '1. Confusion Matrix: Binary Failure Detection',
            'caption': 'Shows 35% of failures missed (15 false negatives). High accuracy (97.9%) misleading due to class imbalance.'
        },
        {
            'name': 'phase_comparison.png',
            'title': '2. Phase Comparison: Model Evolution',
            'caption': 'Phases 1-3 stable around 0.55 F1. Phase 4 catastrophic regression (-41%) due to SMOTE incompatibility.'
        },
        {
            'name': 'training_curves.png',
            'title': '3. Training Loss Curves: Convergence Analysis',
            'caption': 'Model converged at epoch 30. Plateau indicates data limitation, not training issue. Cannot improve further.'
        },
        {
            'name': 'failure_types.png',
            'title': '4. Failure Type Performance: Per-Class Analysis',
            'caption': 'Rare failure types (TWF=5, RNF=3 samples) completely missed (F1=0). Need 10-50× more samples per type.'
        },
        {
            'name': 'roc_curve.png',
            'title': '5. ROC Curve: Evaluation Metric Comparison',
            'caption': 'AUC=0.9612 appears excellent but misleading for imbalanced data. Precision-recall curve is better metric.'
        }
    ]
    
    for viz in visualizations:
        viz_path = Path('visualizations') / viz['name']
        
        if viz_path.exists():
            # Add visualization heading
            viz_subheading = doc.add_heading(viz['title'], level=2)
            viz_subheading.runs[0].font.size = Pt(13)
            
            # Add image
            doc.add_picture(str(viz_path), width=Inches(5.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add caption
            caption_para = doc.add_paragraph(viz['caption'], style='Caption')
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_para.runs[0].font.italic = True
            caption_para.runs[0].font.size = Pt(10)
            
            doc.add_paragraph()  # Spacing
        else:
            print(f"   ⚠ Warning: {viz_path} not found. Generate with generate_visualizations.py first")
            doc.add_heading(viz['title'], level=2)
            doc.add_paragraph(f"[IMAGE NOT FOUND: {viz['name']} - Run generate_visualizations.py first]")
    
    # ========================================================================
    # RESULTS TABLE
    # ========================================================================
    results_heading = doc.add_heading('Detailed Results Summary', level=1)
    results_heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)
    
    # Create results table
    table = doc.add_table(rows=7, cols=4)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Metric'
    header_cells[1].text = 'Expected Target'
    header_cells[2].text = 'Actual Result'
    header_cells[3].text = 'Status'
    
    # Make header bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Data rows
    results_data = [
        ('Binary Failure F1', '>0.95', '0.5545', '❌ MISS'),
        ('Type Classification Accuracy', '>95%', '97.25%', '✅ PASS'),
        ('Time-to-Failure MAE', '<2.0 hrs', '3.50 hrs', '❌ MISS'),
        ('Model Speed', '<50ms', '<10ms', '✅ PASS'),
        ('ROC AUC Score', '>0.98', '0.9612', '⚠ PARTIAL'),
        ('Precision @ Optimal Threshold', '>90%', '48.3%', '❌ MISS'),
    ]
    
    for i, (metric, target, actual, status) in enumerate(results_data, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = metric
        row_cells[1].text = target
        row_cells[2].text = actual
        row_cells[3].text = status
    
    doc.add_paragraph()
    
    # ========================================================================
    # LIMITATIONS SECTION
    # ========================================================================
    lim_heading = doc.add_heading('8 Critical Limitations', level=1)
    lim_heading.runs[0].font.color.rgb = RGBColor(204, 0, 0)
    
    limitations = [
        "Pseudo-temporal data: Dataset has 10,000 static snapshots sorted by tool wear, NOT real time-series. TCN expected temporal dependencies that don't exist.",
        "Extreme class imbalance: 97:3 ratio with only 43 test failures creates ±0.15 confidence interval – statistically impossible to prove improvement.",
        "Synthetic TTF labels: Circular dependency where TTF created from same features used for prediction, making noise floor ~3-4 hours.",
        "Insufficient diversity: Rare failure types (TWF=5, RNF=3 samples) completely missed. Need 50-200× more samples per type.",
        "Multi-task interference: Sharing 133K parameters across 3 tasks forced compromises on each. Single-task models likely better.",
        "Unrealistic targets: F1>0.95 with 43 test failures was never achievable – set based on theory, not data analysis.",
        "Feature artifacts: Temporal features created from sort order rather than real time patterns. Window-based sequences provided zero benefit.",
        "Local optimum: Model converged at epoch 30 with validation loss floor ~0.20. Data ceiling prevents further improvement.",
    ]
    
    for i, limitation in enumerate(limitations, 1):
        doc.add_paragraph(f"{i}. {limitation}", style='List Number')
    
    doc.add_page_break()
    
    # ========================================================================
    # RECOMMENDATIONS
    # ========================================================================
    rec_heading = doc.add_heading('Actionable Recommendations', level=1)
    rec_heading.runs[0].font.color.rgb = RGBColor(0, 153, 0)
    
    # Immediate actions
    doc.add_heading('Immediate Actions (Week 1)', level=2)
    immediate = [
        "Revert Phase 4 implementation – SMOTE + class-balanced batching incompatible",
        "Deploy Phase 3 model with threshold=0.30 (prioritize recall over precision for safety)",
        "Train 5 separate models with different random seeds, average predictions (ensemble)",
        "Use model only as human-in-the-loop system (operator review required)"
    ]
    for item in immediate:
        doc.add_paragraph(item, style='List Bullet')
    
    # Medium-term
    doc.add_heading('Medium-Term Improvements (3-6 months)', level=2)
    medium = [
        "Collect real time-series data from production (not synthetic)",
        "Ground truth TTF labels from actual maintenance records",
        "Accumulate 200+ test failures per type for statistical power",
        "Explore alternative approaches: anomaly detection, survival analysis models"
    ]
    for item in medium:
        doc.add_paragraph(item, style='List Bullet')
    
    # Long-term
    doc.add_heading('Long-Term Strategy (6-12 months)', level=2)
    longterm = [
        "Establish continuous data pipeline for ongoing model improvement",
        "Implement online learning for concept drift adaptation",
        "Consider domain-specific architectures (physics-informed neural networks)",
        "Target realistic F1=0.70-0.80 with 500+ test failures per class"
    ]
    for item in longterm:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # ========================================================================
    # LESSONS LEARNED
    # ========================================================================
    lessons_heading = doc.add_heading('Key Lessons & Project Reflections', level=1)
    lessons_heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)
    
    lessons_text = """
1. DATA QUALITY TRUMPS MODEL COMPLEXITY
   Our advanced TCN with attention mechanisms couldn't overcome fundamentally pseudo-temporal data.
   A simple logistic regression with better data would likely outperform our model.
   Lesson: Always analyze data first, build simple baseline before complex architectures.

2. SYNTHETIC LABELS HAVE HARD LIMITS
   Time-to-failure prediction hit a noise floor at 3-4 hours because features were used to create
   labels. This creates circular dependency making true timing unpredictable.
   Lesson: Validate label generation methodology independently before investing in prediction tasks.

3. CLASS IMBALANCE REQUIRES RETHINKING ENTIRE APPROACH
   97:3 ratio with only 43 test failures made statistical confidence impossible. Focal loss,
   SMOTE, and weighted sampling all insufficient against fundamental data shortage.
   Lesson: Set targets based on data analysis, not domain best practices.

4. MULTI-TASK LEARNING CREATES COMPROMISES
   Sharing representations across 3 unrelated tasks forced the model to find suboptimal
   solutions for each. Single-task models would likely perform better on failure detection.
   Lesson: Multi-task learning helps regularization but can hurt task-specific performance.

5. TEMPORAL MODELS NEED REAL TEMPORAL DATA
   TCN's strength is capturing long-range dependencies. Our window-based sequences from
   sorted data provided zero benefit over single timesteps. Wasted 2 weeks on windowing.
   Lesson: Verify architectural assumptions match data characteristics before implementation.

6. AGGRESSIVE TARGETS WITHOUT DATA ANALYSIS ARE COUNTERPRODUCTIVE
   F1>0.95 with 43 test failures is ±0.15 confidence interval – statistically impossible.
   This unrealistic target didn't drive innovation, just led to Phase 4 desperation (SMOTE).
   Lesson: Set targets after exploratory analysis, not before.

7. ENSEMBLE APPROACHES HAVE DIMINISHING RETURNS
   Phase 1-3 improvements: +0.3%, -0.2%, total +0.1% despite 11 major techniques.
   Suggests ceiling effect more than technique effectiveness.
   Lesson: Validate hypothesis that changes help before heavy investment.

8. COMPREHENSIVE DOCUMENTATION PREVENTS WASTED CYCLES
   Phase 4 SMOTE failure could have been predicted with better documentation of
   what didn't work. Explicit reasoning about why prevented duplicate attempts.
   Lesson: Document failed experiments as thoroughly as successful ones.
    """
    
    doc.add_paragraph(lessons_text)
    
    # ========================================================================
    # FINAL SUMMARY
    # ========================================================================
    final_heading = doc.add_heading('Final Verdict', level=1)
    final_heading.runs[0].font.color.rgb = RGBColor(204, 0, 0)
    
    final_text = """
PERFORMANCE: Targets NOT met (F1: 0.5545 vs >0.95 target)
    Status: ❌ MISS
    Gap: -37%
    Root Cause: Data limitations (97:3 imbalance, pseudo-temporal sequences, synthetic labels)

DEPLOYMENT READINESS: NOT safe for autonomous safety-critical systems
    Status: ❌ UNSAFE SOLO
    Recommendation: Use as human-in-the-loop system only (operator review required)
    Deployment Threshold: 0.30 (prioritize catching failures over false alarms)

PATH FORWARD: 6-12 months to achieve target performance with proper data
    Step 1 (Months 0-3): Collect real time-series data, ground truth labels
    Step 2 (Months 3-6): Accumulate 200+ test failures per class
    Step 3 (Months 6-12): Retrain models with proper data, expect F1=0.70-0.80

ESTIMATED EFFORT: 6-12 person-months total (data collection is bottleneck)

CONFIDENCE LEVEL: MODERATE
    The approach and architecture are sound. With proper data, we can achieve target performance.
    Current results represent best possible with existing dataset constraints.
    """
    
    doc.add_paragraph(final_text)
    
    # ========================================================================
    # SAVE DOCUMENT
    # ========================================================================
    output_path = 'finalreport_with_figures.docx'
    doc.save(output_path)
    
    print(f"\n✅ DOCX created successfully!")
    print(f"   📄 Output: {output_path}")
    print(f"   📊 Embedded {len([v for v in visualizations if (Path('visualizations') / v['name']).exists()])} visualizations")
    
    return output_path


# ============================================================================
# MAIN EXECUTION
# ============================================================================
def main():
    """Main entry point"""
    print("=" * 70)
    print("GENERATING FINAL REPORT WITH EMBEDDED VISUALIZATIONS")
    print("=" * 70)
    
    # Check if visualizations exist
    viz_dir = Path('visualizations')
    if not viz_dir.exists():
        print("\n⚠️  Visualizations directory not found!")
        print("   Run 'python generate_visualizations.py' first to generate PNGs")
        print("   Continuing without images...")
    
    # Create DOCX
    output_path = create_docx_with_figures()
    
    # Verify output
    if Path(output_path).exists():
        file_size = Path(output_path).stat().st_size / (1024 * 1024)
        print(f"   📦 File size: {file_size:.2f} MB")
        print(f"\n✨ Report ready for submission!")
        print(f"   Open with: Microsoft Word, Google Docs, or LibreOffice")
    else:
        print("\n❌ Error: DOCX file not created")


if __name__ == '__main__':
    main()
