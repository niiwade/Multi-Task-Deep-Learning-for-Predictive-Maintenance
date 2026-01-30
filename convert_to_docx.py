#!/usr/bin/env python3
"""
Convert finalreport.md to DOCX format with professional formatting
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def add_horizontal_line(paragraph):
    """Add a horizontal line to a paragraph"""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

def read_markdown():
    """Read the markdown file"""
    with open('finalreport.md', 'r', encoding='utf-8') as f:
        return f.read()

def create_docx():
    """Create DOCX document from markdown"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Read markdown
    content = read_markdown()
    lines = content.split('\n')
    
    # Track state
    in_code_block = False
    in_table = False
    table_rows = []
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Title (# ... level 1)
        if line.startswith('# ') and not line.startswith('## '):
            title = line[2:].strip()
            heading = doc.add_heading(title, level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in heading.runs:
                run.font.size = Pt(16)
                run.font.bold = True
            i += 1
            continue
        
        # Heading 2 (## ...)
        if line.startswith('## '):
            heading = doc.add_heading(line[3:].strip(), level=1)
            for run in heading.runs:
                run.font.size = Pt(14)
                run.font.bold = True
            i += 1
            continue
        
        # Heading 3 (### ...)
        if line.startswith('### '):
            heading = doc.add_heading(line[4:].strip(), level=2)
            for run in heading.runs:
                run.font.size = Pt(12)
                run.font.bold = True
            i += 1
            continue
        
        # Heading 4 (#### ...)
        if line.startswith('#### '):
            heading = doc.add_heading(line[5:].strip(), level=3)
            for run in heading.runs:
                run.font.size = Pt(11)
                run.font.bold = True
            i += 1
            continue
        
        # Horizontal rule
        if line.startswith('---'):
            p = doc.add_paragraph()
            add_horizontal_line(p)
            i += 1
            continue
        
        # Bullet points
        if line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
            i += 1
            continue
        
        # Numbered list
        if re.match(r'^\d+\.\s', line):
            match = re.match(r'^(\d+)\.\s(.+)', line)
            if match:
                num, text = match.groups()
                doc.add_paragraph(text, style='List Number')
            i += 1
            continue
        
        # Code block
        if line.startswith('```'):
            if in_code_block:
                in_code_block = False
            else:
                in_code_block = True
                # Add code block content
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                
                code_text = '\n'.join(code_lines).strip()
                if code_text:
                    p = doc.add_paragraph(code_text, style='List Paragraph')
                    for run in p.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                    p.paragraph_format.left_indent = Inches(0.5)
                continue
        
        # Table (markdown format with |)
        if '|' in line and not in_code_block:
            # Collect table rows
            table_rows = [line]
            i += 1
            
            # Get separator row
            if i < len(lines) and '|' in lines[i] and '-' in lines[i]:
                i += 1  # Skip separator
            
            # Get remaining rows
            while i < len(lines) and '|' in lines[i] and not lines[i].startswith('#'):
                table_rows.append(lines[i].strip())
                i += 1
            
            if table_rows:
                # Parse and create table
                cells_data = []
                for row in table_rows:
                    cells = [cell.strip() for cell in row.split('|')[1:-1]]  # Remove empty first/last
                    cells_data.append(cells)
                
                if cells_data:
                    num_cols = len(cells_data[0])
                    num_rows = len(cells_data)
                    
                    table = doc.add_table(rows=num_rows, cols=num_cols)
                    table.style = 'Light Grid Accent 1'
                    
                    for row_idx, row_data in enumerate(cells_data):
                        for col_idx, cell_text in enumerate(row_data):
                            cell = table.rows[row_idx].cells[col_idx]
                            cell_para = cell.paragraphs[0]
                            cell_para.text = cell_text
                            
                            # Bold header row
                            if row_idx == 0:
                                for run in cell_para.runs:
                                    run.font.bold = True
            continue
        
        # Bold and italic formatting
        # Regular paragraph
        if not line.startswith('|'):
            paragraph = doc.add_paragraph()
            
            # Process inline formatting
            text = line
            current_pos = 0
            
            while current_pos < len(text):
                # Find next formatting marker
                bold_pos = text.find('**', current_pos)
                italic_pos = text.find('*', current_pos)
                code_pos = text.find('`', current_pos)
                
                positions = []
                if bold_pos != -1:
                    positions.append(('**', bold_pos))
                if italic_pos != -1 and italic_pos != bold_pos:
                    positions.append(('*', italic_pos))
                if code_pos != -1:
                    positions.append(('`', code_pos))
                
                if not positions:
                    # No formatting, add rest of text
                    if current_pos < len(text):
                        run = paragraph.add_run(text[current_pos:])
                        run.font.name = 'Calibri'
                    break
                
                # Find nearest marker
                marker_type, marker_pos = min(positions, key=lambda x: x[1])
                
                # Add text before marker
                if marker_pos > current_pos:
                    run = paragraph.add_run(text[current_pos:marker_pos])
                    run.font.name = 'Calibri'
                
                # Handle formatting
                if marker_type == '**':
                    # Bold
                    end_pos = text.find('**', marker_pos + 2)
                    if end_pos != -1:
                        bold_text = text[marker_pos+2:end_pos]
                        run = paragraph.add_run(bold_text)
                        run.font.bold = True
                        run.font.name = 'Calibri'
                        current_pos = end_pos + 2
                    else:
                        current_pos = marker_pos + 2
                
                elif marker_type == '*':
                    # Italic
                    end_pos = text.find('*', marker_pos + 1)
                    if end_pos != -1 and text[end_pos-1:end_pos+1] != '**':
                        italic_text = text[marker_pos+1:end_pos]
                        run = paragraph.add_run(italic_text)
                        run.italic = True
                        run.font.name = 'Calibri'
                        current_pos = end_pos + 1
                    else:
                        current_pos = marker_pos + 1
                
                elif marker_type == '`':
                    # Code
                    end_pos = text.find('`', marker_pos + 1)
                    if end_pos != -1:
                        code_text = text[marker_pos+1:end_pos]
                        run = paragraph.add_run(code_text)
                        run.font.name = 'Courier New'
                        run.font.size = Pt(10)
                        current_pos = end_pos + 1
                    else:
                        current_pos = marker_pos + 1
        
        i += 1
    
    return doc

def main():
    """Main conversion function"""
    print("Converting finalreport.md to DOCX format...")
    
    doc = create_docx()
    
    # Save document
    output_file = 'finalreport.docx'
    doc.save(output_file)
    
    print(f"✅ Successfully created: {output_file}")
    print(f"Location: {output_file}")

if __name__ == '__main__':
    main()
